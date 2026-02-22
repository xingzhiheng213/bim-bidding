"""Markdown to DOCX conversion for stage 4.3.

Uses markdown + beautifulsoup4 to parse Markdown, then python-docx to build the document.
Handles # ## ### headings, - lists, paragraphs, **bold**, *italic*, tables (optional).
Stage 7.2: supports format_options for heading/body/table font, first_line_indent, line_spacing.
Word OOXML: font.name sets w:ascii/w:hAnsi (Latin); Chinese uses w:eastAsia, so we set both.
"""
from __future__ import annotations

import re
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt

from app.settings_store import DEFAULT_EXPORT_FORMAT

# Cover page title default font size (pt)
COVER_TITLE_SIZE_PT = 26
# TOC: indent per level (pt) - only used when not using TOC field
TOC_INDENT_PT_PER_LEVEL = 24
# Word OOXML namespace for parse_xml
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _set_run_font_name(run, font_name: str) -> None:
    """Set run font name for both Latin (ascii/hAnsi) and East Asian (eastAsia) so 中文 uses the same font."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is not None:
        rFonts.set(qn("w:eastAsia"), font_name)
    else:
        # Some python-docx versions don't create rFonts when setting name; create it for eastAsia
        rFonts = parse_xml(
            f'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:eastAsia="{font_name}"/>'
        )
        rPr.insert(0, rFonts)


def _add_inline_to_paragraph(p, elem) -> None:
    """Add inline content (bold, italic, etc) from elem to paragraph p."""
    for child in elem:
        if hasattr(child, "name") and child.name:
            if child.name in ("strong", "b"):
                r = p.add_run(child.get_text())
                r.bold = True
            elif child.name in ("em", "i"):
                r = p.add_run(child.get_text())
                r.italic = True
            elif child.name == "code":
                p.add_run(child.get_text())
            elif child.name == "br":
                p.add_run("\n")
            else:
                p.add_run(child.get_text() if hasattr(child, "get_text") else str(child))
        elif isinstance(child, str) or (hasattr(child, "name") and child.name is None):
            txt = str(child) if child else ""
            if txt:
                p.add_run(txt)


def _apply_heading_style(paragraph, level: int, opts: dict) -> None:
    """Apply heading font/size from opts (heading_1/2/3) to paragraph runs; level 4-6 use heading_3."""
    lvl = min(level, 3)
    font_key = f"heading_{lvl}_font"
    size_key = f"heading_{lvl}_size_pt"
    font_name = opts.get(font_key)
    size_pt = opts.get(size_key)
    if font_name is None and size_pt is None:
        return
    for run in paragraph.runs:
        if font_name is not None:
            _set_run_font_name(run, font_name)
        if size_pt is not None:
            run.font.size = Pt(size_pt)


def _apply_body_style(paragraph, opts: dict) -> None:
    """Apply body font/size, first_line_indent, line_spacing to paragraph."""
    font_name = opts.get("body_font")
    size_pt = opts.get("body_size_pt")
    for run in paragraph.runs:
        if font_name is not None:
            _set_run_font_name(run, font_name)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
    indent_pt = opts.get("first_line_indent_pt")
    if indent_pt is not None:
        paragraph.paragraph_format.first_line_indent = Pt(indent_pt)
    spacing = opts.get("line_spacing")
    if spacing is not None:
        paragraph.paragraph_format.line_spacing = spacing


def _apply_table_cell_style(cell, opts: dict) -> None:
    """Apply table font/size to first paragraph runs in cell."""
    font_name = opts.get("table_font")
    size_pt = opts.get("table_size_pt")
    if font_name is None and size_pt is None:
        return
    if not cell.paragraphs:
        return
    for run in cell.paragraphs[0].runs:
        if font_name is not None:
            _set_run_font_name(run, font_name)
        if size_pt is not None:
            run.font.size = Pt(size_pt)


def _add_table_from_cell_texts(doc: Document, cell_texts: list[list[str]], opts: dict) -> None:
    """Append a Word table to doc from cell_texts (list of rows, each row list of cell strings)."""
    if not cell_texts:
        return
    max_cols = max(len(r) for r in cell_texts)
    for r in cell_texts:
        while len(r) < max_cols:
            r.append("")
    table = doc.add_table(rows=len(cell_texts), cols=max_cols)
    table.style = "Table Grid"
    for i, row_texts in enumerate(cell_texts):
        for j, text in enumerate(row_texts):
            table.rows[i].cells[j].text = text
            _apply_table_cell_style(table.rows[i].cells[j], opts)


def _add_table_from_soup(doc: Document, table_elem, opts: dict) -> None:
    """Parse a BeautifulSoup <table> element and append a Word table to doc."""
    rows = table_elem.find_all("tr")
    if not rows:
        return
    cell_texts = []
    for tr in rows:
        cells = tr.find_all(["th", "td"])
        cell_texts.append([c.get_text(strip=True) for c in cells])
    _add_table_from_cell_texts(doc, cell_texts, opts)


def _is_separator_row(cells: list[str]) -> bool:
    """True if all cells look like markdown table separator (e.g. --- or :---)."""
    if not cells:
        return False
    return all(c.strip().replace(":", "").replace("-", "").strip() == "" for c in cells)


def _looks_like_table_row(text: str) -> bool:
    """True if text looks like a markdown table row (starts with | and contains at least one more |)."""
    if not text or not text.strip():
        return False
    t = text.strip()
    return t.startswith("|") and "|" in t[1:]

def _split_md_table_cells(line: str) -> list[str]:
    """Split a markdown pipe-table row into cells (strip outer pipes and whitespace)."""
    cells = [c.strip() for c in line.strip().split("|")]
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def _is_md_table_separator_line(line: str) -> bool:
    """Return True if line looks like a markdown table separator row."""
    if not line or "|" not in line:
        return False
    cells = _split_md_table_cells(line)
    return _is_separator_row(cells)


def _preprocess_markdown_tables(markdown_text: str) -> str:
    """Normalize markdown pipe tables to make parsers robust.

    Key goals:
    - Only treat a region as a table if it has a header row AND a separator row (--- / :--- etc).
      This avoids misclassifying normal text with '|' as a table.
    - Remove blank lines INSIDE the table (between header/separator/body rows), as blank lines
      terminate tables in common markdown parsers (Pandoc/GFM/Python-Markdown).
    - Ensure there is a blank line BEFORE a table block when it directly follows text.
    - Do NOT touch anything inside fenced code blocks (```), where '|' may appear in examples.
    """
    if not markdown_text:
        return markdown_text

    lines = markdown_text.splitlines()
    out: list[str] = []
    i = 0
    in_fence = False

    def _is_fence(line: str) -> bool:
        s = line.lstrip()
        return s.startswith("```")

    def _is_md_table_row_line(line: str) -> bool:
        s = line.strip()
        # Stronger than just containing '|': require starting with '|' and at least 2 pipes total.
        return s.startswith("|") and s.count("|") >= 2

    while i < len(lines):
        line = lines[i]

        # Track fenced code blocks; do not preprocess inside.
        if _is_fence(line):
            in_fence = not in_fence
            out.append(line)
            i += 1
            continue
        if in_fence:
            out.append(line)
            i += 1
            continue

        # Try to detect a table starting at i: header row + separator row (allow blank lines between).
        if _is_md_table_row_line(line):
            header_cells = _split_md_table_cells(line)
            # Need at least 2 columns to be a table.
            if len(header_cells) >= 2:
                j = i + 1
                # Allow blank lines between header and separator (common "bad" output).
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if j < len(lines) and _is_md_table_separator_line(lines[j]):
                    sep_cells = _split_md_table_cells(lines[j])
                    # Separator column count must match header.
                    if len(sep_cells) == len(header_cells):
                        # Ensure blank line BEFORE table (if previous output line is non-empty)
                        if out and out[-1].strip():
                            out.append("")

                        # Emit header + separator
                        out.append("| " + " | ".join(header_cells) + " |")
                        out.append("|" + "|".join(["---"] * len(header_cells)) + "|")

                        # Emit body rows: collect subsequent table rows; ignore blank lines between.
                        k = j + 1
                        while k < len(lines):
                            if _is_fence(lines[k]):
                                break
                            if not lines[k].strip():
                                k += 1
                                continue
                            if not _is_md_table_row_line(lines[k]):
                                break
                            row_cells = _split_md_table_cells(lines[k])
                            # Keep only rows with matching column count; otherwise stop (avoid consuming text).
                            if len(row_cells) != len(header_cells):
                                break
                            out.append("| " + " | ".join(row_cells) + " |")
                            k += 1

                        # Ensure a blank line AFTER table if next output would stick to it.
                        if k < len(lines) and lines[k].strip():
                            out.append("")

                        i = k
                        continue

        # Default: passthrough line as-is
        out.append(line)
        i += 1

    return "\n".join(out)


def _strip_bold_in_string(s: str) -> str:
    """Remove Markdown bold (**...** and __...__) from a string. Repeats until no change (handles adjacent/nested)."""
    if not s:
        return s
    prev = None
    while prev != s:
        prev = s
        s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
        s = re.sub(r"__(.+?)__", r"\1", s)
    return s


def _get_table_line_indices(lines: list[str]) -> set[int]:
    """Return set of line indices that belong to a pipe table block (header + separator + body).
    Used so we do NOT strip bold inside table cells. Fenced code blocks are not treated as tables.
    """
    table_indices: set[int] = set()
    i = 0
    in_fence = False

    def _is_fence(line: str) -> bool:
        return line.lstrip().startswith("```")

    def _is_md_table_row_line(line: str) -> bool:
        t = line.strip()
        return t.startswith("|") and t.count("|") >= 2

    while i < len(lines):
        line = lines[i]
        if _is_fence(line):
            in_fence = not in_fence
            i += 1
            continue
        if in_fence:
            i += 1
            continue
        if _is_md_table_row_line(line):
            header_cells = _split_md_table_cells(line)
            if len(header_cells) >= 2:
                j = i + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if j < len(lines) and _is_md_table_separator_line(lines[j]):
                    sep_cells = _split_md_table_cells(lines[j])
                    if len(sep_cells) == len(header_cells):
                        table_indices.add(i)
                        table_indices.add(j)
                        k = j + 1
                        while k < len(lines):
                            if _is_fence(lines[k]):
                                break
                            if not lines[k].strip():
                                k += 1
                                continue
                            if not _is_md_table_row_line(lines[k]):
                                break
                            row_cells = _split_md_table_cells(lines[k])
                            if len(row_cells) != len(header_cells):
                                break
                            table_indices.add(k)
                            k += 1
                        i = k
                        continue
        i += 1
    return table_indices


def _strip_bold_outside_tables(markdown_text: str) -> str:
    """Remove ** and __ bold from lines that are NOT inside a pipe table block. Keeps bold in table cells (e.g. 合计)."""
    if not markdown_text:
        return markdown_text
    lines = markdown_text.splitlines()
    table_indices = _get_table_line_indices(lines)
    out = []
    for i, line in enumerate(lines):
        if i in table_indices:
            out.append(line)
        else:
            out.append(_strip_bold_in_string(line))
    return "\n".join(out)


def _extract_scoring_comment(text: str) -> str | None:
    """If paragraph is '（评分响应：...）' or '（（评分响应：...））', return the comment text; else None."""
    if not text or not isinstance(text, str):
        return None
    s = text.strip()
    while len(s) >= 2 and s[0] in "（(" and s[-1] in "）)":
        s = s[1:-1].strip()
    if s.startswith("评分响应：") or s.startswith("评分响应:"):
        return s.strip()
    return None


def _apply_pending_comment(doc: Document, p, pending_comment: str | None) -> bool:
    """If pending_comment is set and p has runs, add it as a Word comment to p. Returns True if comment was added."""
    if not pending_comment or not p or not getattr(p, "runs", None):
        return False
    if list(p.runs):
        doc.add_comment(p.runs, text=pending_comment, author="评分说明", initials="")
        return True
    return False


def _parse_raw_markdown_table(text: str) -> tuple[list[list[str]], str] | None:
    """If text contains raw markdown table (pipes and dashes), return (cell_texts, leading_text); else None.

    Handles tables that were inside code blocks (e.g. with a title line like "表2-1 ..." above).
    leading_text is any content before the first table line (e.g. table caption), for caller to add as paragraph.
    """
    lines = [line.strip() for line in text.strip().splitlines() if line.strip()]
    if len(lines) < 2:
        return None
    start = 0
    for i, line in enumerate(lines):
        if "|" in line:
            start = i
            break
    else:
        return None
    table_lines = lines[start:]
    rows = []
    for line in table_lines:
        if "|" not in line:
            break
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        if not cells:
            continue
        if _is_separator_row(cells):
            continue
        rows.append(cells)
    if len(rows) < 1:
        return None
    leading_text = "\n".join(lines[:start]).strip() if start > 0 else ""
    return (rows, leading_text)


def add_cover_page(doc: Document, title: str, opts: dict) -> None:
    """Add a cover page: one centered title paragraph, then page break."""
    p = doc.add_paragraph(title.strip() or "请输入BIM技术标标题")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        font_name = opts.get("heading_1_font") or opts.get("body_font")
        if font_name:
            _set_run_font_name(run, font_name)
        run.font.size = Pt(opts.get("heading_1_size_pt") or COVER_TITLE_SIZE_PT)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_toc_field_paragraph(doc: Document) -> None:
    """Append a paragraph containing a Word TOC field. Word/WPS will show placeholder until user updates the field."""
    p = doc.add_paragraph()
    # TOC field: begin, instrText, separate, end (each in its own w:r)
    p._element.append(
        parse_xml(
            f'<w:r xmlns:w="{W_NS}"><w:fldChar w:fldCharType="begin"/></w:r>'
        )
    )
    p._element.append(
        parse_xml(
            f'<w:r xmlns:w="{W_NS}"><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
        )
    )
    p._element.append(
        parse_xml(
            f'<w:r xmlns:w="{W_NS}"><w:fldChar w:fldCharType="separate"/></w:r>'
        )
    )
    p._element.append(
        parse_xml(
            f'<w:r xmlns:w="{W_NS}"><w:fldChar w:fldCharType="end"/></w:r>'
        )
    )


def add_toc_page(doc: Document, toc_entries: list[tuple[int, str]], opts: dict) -> None:
    """Add a TOC page: '目录' heading, then a TOC field (Word/WPS 更新域后可生成三级目录), then page break."""
    p = doc.add_heading("目录", level=1)
    _apply_heading_style(p, 1, opts)
    _add_toc_field_paragraph(doc)
    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)


def markdown_to_docx(
    markdown_text: str,
    format_options: dict | None = None,
    doc: Document | None = None,
) -> Document:
    """Convert Markdown string to python-docx Document (or append to existing doc).

    Handles # ## ### headings, - lists, paragraphs, **bold**, *italic*, tables.
    format_options: optional dict from get_export_format_config(); merged with DEFAULT_EXPORT_FORMAT.
    doc: if provided, content is appended to this document; otherwise a new Document is created.
    Returns the Document instance (call .save() to write to file).
    """
    opts = {**DEFAULT_EXPORT_FORMAT, **(format_options or {})}
    if doc is None:
        doc = Document()

    if not markdown_text or not markdown_text.strip():
        p = doc.add_paragraph("（无内容）")
        _apply_body_style(p, opts)
        return doc

    # Preprocess markdown tables to avoid blank-line table breakage.
    markdown_text = _preprocess_markdown_tables(markdown_text)
    # Strip bold in body/headings; keep bold only inside table cells.
    markdown_text = _strip_bold_outside_tables(markdown_text)

    html = markdown.markdown(
        markdown_text,
        extensions=["tables", "fenced_code", "nl2br"],
    )
    soup = BeautifulSoup(html, "html.parser")
    root = soup.find("body") or soup
    children = [c for c in root.children if hasattr(c, "name") and c.name is not None]
    skip_until = -1
    last_paragraph = None
    pending_comment = None

    for i, elem in enumerate(children):
        if i <= skip_until:
            continue
        if not hasattr(elem, "name") or elem.name is None:
            continue
        if elem.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(elem.name[1])
            p = doc.add_heading(elem.get_text(strip=True), level=level)
            _apply_heading_style(p, level, opts)
            if _apply_pending_comment(doc, p, pending_comment):
                pending_comment = None
            last_paragraph = p
        elif elem.name == "p":
            text = elem.get_text(strip=True)
            comment_text = _extract_scoring_comment(text)
            if comment_text is not None:
                if last_paragraph and list(getattr(last_paragraph, "runs", [])):
                    doc.add_comment(last_paragraph.runs, text=comment_text, author="评分说明", initials="")
                else:
                    pending_comment = comment_text
                continue
            if _looks_like_table_row(text):
                lines = [text]
                j = i + 1
                while j < len(children):
                    next_elem = children[j]
                    if getattr(next_elem, "name", None) != "p":
                        break
                    next_text = next_elem.get_text(strip=True)
                    if not next_text:
                        j += 1
                        continue
                    if not _looks_like_table_row(next_text):
                        break
                    lines.append(next_text)
                    j += 1
                block = "\n".join(lines)
                parsed = _parse_raw_markdown_table(block)
                if parsed:
                    cell_texts, leading_text = parsed
                    if leading_text:
                        p = doc.add_paragraph(leading_text)
                        _apply_body_style(p, opts)
                        if _apply_pending_comment(doc, p, pending_comment):
                            pending_comment = None
                        last_paragraph = p
                    _add_table_from_cell_texts(doc, cell_texts, opts)
                    skip_until = j - 1
                    continue
            if text:
                p = doc.add_paragraph()
                _add_inline_to_paragraph(p, elem)
            else:
                p = doc.add_paragraph()
            _apply_body_style(p, opts)
            if _apply_pending_comment(doc, p, pending_comment):
                pending_comment = None
            last_paragraph = p
        elif elem.name in ("ul", "ol"):
            list_style = "List Bullet" if elem.name == "ul" else "List Number"
            for li in elem.find_all("li", recursive=False):
                text = li.get_text(strip=True)
                if text:
                    p = doc.add_paragraph(style=list_style)
                    _add_inline_to_paragraph(p, li)
                else:
                    p = doc.add_paragraph(style=list_style)
                _apply_body_style(p, opts)
                if _apply_pending_comment(doc, p, pending_comment):
                    pending_comment = None
                last_paragraph = p
        elif elem.name == "table":
            _add_table_from_soup(doc, elem, opts)
        elif elem.name == "hr":
            p = doc.add_paragraph()
            _apply_body_style(p, opts)
            if _apply_pending_comment(doc, p, pending_comment):
                pending_comment = None
            last_paragraph = p
        elif elem.name == "pre":
            pre_text = elem.get_text()
            parsed = _parse_raw_markdown_table(pre_text)
            if parsed:
                cell_texts, leading_text = parsed
                if leading_text:
                    p = doc.add_paragraph(leading_text)
                    _apply_body_style(p, opts)
                    if _apply_pending_comment(doc, p, pending_comment):
                        pending_comment = None
                    last_paragraph = p
                _add_table_from_cell_texts(doc, cell_texts, opts)
            else:
                p = doc.add_paragraph(pre_text)
                _apply_body_style(p, opts)
                if _apply_pending_comment(doc, p, pending_comment):
                    pending_comment = None
                last_paragraph = p
        elif elem.name == "div":
            div_children = [c for c in elem.children if hasattr(c, "name") and c.name]
            div_skip_until = -1
            for idx, sub in enumerate(div_children):
                if idx <= div_skip_until:
                    continue
                if sub.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                    level = int(sub.name[1])
                    p = doc.add_heading(sub.get_text(strip=True), level=level)
                    _apply_heading_style(p, level, opts)
                    if _apply_pending_comment(doc, p, pending_comment):
                        pending_comment = None
                    last_paragraph = p
                elif sub.name == "p":
                    text = sub.get_text(strip=True)
                    comment_text = _extract_scoring_comment(text)
                    if comment_text is not None:
                        if last_paragraph and list(getattr(last_paragraph, "runs", [])):
                            doc.add_comment(last_paragraph.runs, text=comment_text, author="评分说明", initials="")
                        else:
                            pending_comment = comment_text
                        continue
                    if _looks_like_table_row(text):
                        lines = [text]
                        k = idx + 1
                        while k < len(div_children):
                            next_sub = div_children[k]
                            if getattr(next_sub, "name", None) != "p":
                                break
                            next_text = next_sub.get_text(strip=True)
                            if not next_text:
                                k += 1
                                continue
                            if not _looks_like_table_row(next_text):
                                break
                            lines.append(next_text)
                            k += 1
                        block = "\n".join(lines)
                        parsed = _parse_raw_markdown_table(block)
                        if parsed:
                            cell_texts, leading_text = parsed
                            if leading_text:
                                p = doc.add_paragraph(leading_text)
                                _apply_body_style(p, opts)
                                if _apply_pending_comment(doc, p, pending_comment):
                                    pending_comment = None
                                last_paragraph = p
                            _add_table_from_cell_texts(doc, cell_texts, opts)
                            div_skip_until = k - 1
                            continue
                    if text:
                        p = doc.add_paragraph()
                        _add_inline_to_paragraph(p, sub)
                    else:
                        p = doc.add_paragraph()
                    _apply_body_style(p, opts)
                    if _apply_pending_comment(doc, p, pending_comment):
                        pending_comment = None
                    last_paragraph = p
                elif sub.name in ("ul", "ol"):
                    list_style = "List Bullet" if sub.name == "ul" else "List Number"
                    for li in sub.find_all("li", recursive=False):
                        p = doc.add_paragraph(li.get_text(strip=True), style=list_style)
                        _apply_body_style(p, opts)
                        if _apply_pending_comment(doc, p, pending_comment):
                            pending_comment = None
                        last_paragraph = p
                elif sub.name == "table":
                    _add_table_from_soup(doc, sub, opts)
                elif sub.name == "pre":
                    pre_text = sub.get_text()
                    parsed = _parse_raw_markdown_table(pre_text)
                    if parsed:
                        cell_texts, leading_text = parsed
                        if leading_text:
                            p = doc.add_paragraph(leading_text)
                            _apply_body_style(p, opts)
                            if _apply_pending_comment(doc, p, pending_comment):
                                pending_comment = None
                            last_paragraph = p
                        _add_table_from_cell_texts(doc, cell_texts, opts)
                    else:
                        p = doc.add_paragraph(pre_text)
                        _apply_body_style(p, opts)
                        if _apply_pending_comment(doc, p, pending_comment):
                            pending_comment = None
                        last_paragraph = p

    return doc
